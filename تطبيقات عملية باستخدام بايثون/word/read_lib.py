import docx
def get_my_text(file_name):

    text_name=docx.Document(file_name)
    all_text=''
    for i in range(len(text_name.paragraphs)):
        x=text_name.paragraphs[i].text
        all_text+=x
    print(all_text)
