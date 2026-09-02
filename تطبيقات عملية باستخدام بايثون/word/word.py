import docx
from pathlib import Path
from read_lib import get_my_text

doc=docx.Document(Path.home()/Path('Desktop','docx_files','Academy_1.docx'))
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[2].text)

print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)
print(doc.paragraphs[2].runs[4].text)




get_my_text(Path.home()/Path('Desktop','docx_files','Academy_1.docx'))
