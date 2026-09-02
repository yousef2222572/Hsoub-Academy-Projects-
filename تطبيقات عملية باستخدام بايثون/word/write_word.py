import docx 
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as al
import docx.shared
doc=docx.Document()
x=docx.Document()
doc.add_paragraph('hello')
doc.add_paragraph('hello i am yosuef i am the rewial one ')


paragraph=doc.add_paragraph('مرحبا بالعالم')
paragraph.alignment=al.RIGHT

paragraph.add_run('hello')



doc.add_heading('title',0)
doc.add_heading('title',1)
doc.add_heading('title',2)
doc.add_heading('title',3)
doc.add_heading('title',4)
doc.add_heading('title',5)
doc.add_heading('title',6)
doc.add_heading('title',7)








#stayle

print(doc.paragraphs[0].style)
print(doc.paragraphs[5].style)
print(doc.paragraphs[3].style)
print(doc.paragraphs[5].text)
print(doc.paragraphs[4].text)

doc.paragraphs[0].style='Title'
doc.paragraphs[0].style=doc.styles['Heading 1']
doc.paragraphs[0].style=doc.paragraphs[2].style
doc.paragraphs[1].style.delete()


doc.add_picture(str(Path.home()/Path('Desktop','Desktop .jpeg')),width=docx.shared.Inches(5),height=docx.shared.Inches(7))























doc.save(Path.home()/Path('Desktop','academy_2.docx'))
